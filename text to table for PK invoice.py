import re
from docx import Document
from docx.shared import Cm, Pt



document_name = '/home/mark/Downloads/Kijuu invoice May-June 2021'
input_document = '/home/mark/Downloads/Kazi ya sil May.txt'




with open(input_document, 'r') as file:
    data = file.readlines()

word_document = Document()
style = word_document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
table = word_document.add_table(0, 0) # we add rows iteratively
table.style = 'TableGrid'
first_column_width = 5
second_column_with = 8
third_column_with = 3
table.add_column(Cm(first_column_width))
table.add_column(Cm(second_column_with))
table.add_column(Cm(third_column_with))
text = list()
total = 0
hours = list()

regex1 = re.compile(r'(\d{1,2}\/\d{1,2}\/?\d{0,4})\s?(.*)')
regex2 = re.compile(r'(.*)\s?([mM]asaa\s?\d{1,3}\.?\s?\d{0,3})')
regex3 = re.compile(r'[mM]asaa\s?(\d{1,3}\.?\s?\d{0,3})')

for i in range(len(data)):
    match1 = regex1.search(data[i])
    if match1 != None:
        string1 = match1.group(1)
        string2 = match1.group(2)
        # text.append(match1.groups())
        match2 = regex2.search(string2)
        if match2 != None:
            string2 = match2.group(1)
            string3 = match2.group(2)
            match3 = regex3.search(string3)
            if match3 != None:
                hours.append(float(match3.group(1)))
                total += hours[len(hours)-1]
        else:
            string3 = ''
        text.append([string1, string2, string3])
    
for index, value in enumerate(text):
    print(value)
        
    table.add_row()
    row = table.rows[index]
    row.cells[0].text = value[0]
    row.cells[1].text = value[1]
    row.cells[2].text = value[2]
    
para1 = word_document.add_paragraph("\nJumla: Masaa " + str(total) + " @ 10,000/= kwa saa\n")  # left align
para2 = word_document.add_paragraph(str(f'{int(total)*10000:,}' + '/='))  # right align
print(f'{int(total)*10000:,}')
para1.style = word_document.styles['Normal']
para1.alignment = 0
para2.style = word_document.styles['Normal']
para2.alignment = 2

table.style = word_document.styles['Table Grid']
table.style.font.name = 'Arial'

word_document.save(document_name + '.docx')


